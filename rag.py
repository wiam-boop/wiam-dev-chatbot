"""
محرك RAG — يستخدم HuggingFace للـ embeddings و Groq (Qwen Vision) لفهم الصور
+ Semantic Cache: يحفظ الإجابات ويسترجعها للأسئلة المتشابهة
"""
import base64
import os
import json
import re
import uuid
import threading

import numpy as np
import faiss
from huggingface_hub import InferenceClient
from groq import Groq

from pypdf import PdfReader
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

# ─── مسار التخزين ───
IS_RAILWAY = os.environ.get("RAILWAY_ENVIRONMENT") is not None

_env_storage_path = os.environ.get("STORAGE_PATH", "").strip()

if _env_storage_path:
    STORAGE_DIR = _env_storage_path
elif IS_RAILWAY:
    STORAGE_DIR = "/tmp/storage"
else:
    STORAGE_DIR = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "storage"
    )

UPLOADS_DIR  = os.path.join(STORAGE_DIR, "uploads")
INDEX_PATH   = os.path.join(STORAGE_DIR, "kb.index")
META_PATH    = os.path.join(STORAGE_DIR, "kb_meta.json")

# ─── مسارات الـ Semantic Cache ───
CACHE_INDEX_PATH = os.path.join(STORAGE_DIR, "cache.index")
CACHE_META_PATH  = os.path.join(STORAGE_DIR, "cache_meta.json")

os.makedirs(UPLOADS_DIR, exist_ok=True)

# ─── HuggingFace embeddings ───
HF_API_KEY       = os.environ.get("HF_API_KEY")
_hf_client       = InferenceClient(provider="hf-inference", api_key=HF_API_KEY)
EMBED_MODEL_NAME = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
EMBED_DIM        = 384

# ─── Groq Vision ───
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
_groq_client = Groq(api_key=GROQ_API_KEY)
VISION_MODEL = "qwen/qwen3.6-27b"

CHUNK_SIZE    = 700
CHUNK_OVERLAP = 120
TOP_K         = 4
_lock         = threading.Lock()

# ─── إعدادات الـ Cache ───
CACHE_SIMILARITY_THRESHOLD = 0.92   # نسبة التشابه المطلوبة (92%)
CACHE_MAX_SIZE             = 500    # أقصى عدد إجابات محفوظة


def _get_embeddings(texts: list) -> np.ndarray:
    vectors = [
        _hf_client.feature_extraction(text, model=EMBED_MODEL_NAME)
        for text in texts
    ]
    embeddings = np.array(vectors, dtype="float32")
    norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
    norms = np.where(norms == 0, 1, norms)
    return embeddings / norms


# =========================================================
# SEMANTIC CACHE
# =========================================================

class SemanticCache:
    """
    ذاكرة ذكية — تحفظ الأسئلة والإجابات وتسترجعها
    عند ورود سؤال مشابه بنسبة تشابه عالية.
    خفيفة على الذاكرة: FAISS + JSON فقط.
    """

    def __init__(self):
        self.index = None
        self.meta  = []   # [{question, answer, hits}, ...]
        self._lock = threading.Lock()
        self._load()

    def _load(self):
        if os.path.exists(CACHE_INDEX_PATH) and os.path.exists(CACHE_META_PATH):
            try:
                self.index = faiss.read_index(CACHE_INDEX_PATH)
                with open(CACHE_META_PATH, "r", encoding="utf-8") as f:
                    self.meta = json.load(f)
                print(f"[CACHE] تم تحميل {len(self.meta)} إجابة محفوظة")
            except Exception as e:
                print(f"[CACHE] خطأ في التحميل: {e}")
                self._reset()
        else:
            self._reset()

    def _reset(self):
        self.index = faiss.IndexFlatIP(EMBED_DIM)
        self.meta  = []

    def _save(self):
        try:
            faiss.write_index(self.index, CACHE_INDEX_PATH)
            with open(CACHE_META_PATH, "w", encoding="utf-8") as f:
                json.dump(self.meta, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[CACHE] خطأ في الحفظ: {e}")

    def get(self, question: str):
        """ابحث عن إجابة مشابهة — يُعيد الإجابة أو None"""
        if self.index.ntotal == 0 or not question.strip():
            return None
        try:
            q_emb           = _get_embeddings([question])
            scores, indices = self.index.search(q_emb, 1)
            score = float(scores[0][0])
            idx   = int(indices[0][0])

            if score >= CACHE_SIMILARITY_THRESHOLD and idx < len(self.meta):
                item = self.meta[idx]
                item["hits"] = item.get("hits", 0) + 1
                self._save()
                print(f"[CACHE HIT] تشابه={score:.2f} | السؤال: {item['question'][:50]}")
                return item["answer"]
        except Exception as e:
            print(f"[CACHE GET ERROR] {e}")
        return None

    def set(self, question: str, answer: str):
        """احفظ سؤالاً وإجابته في الذاكرة"""
        if not question.strip() or not answer.strip():
            return
        # لا تحفظ الأسئلة القصيرة جداً (أقل من 5 أحرف)
        if len(question.strip()) < 5:
            return
        try:
            with self._lock:
                # إذا وصلنا للحد الأقصى، احذف الأقل استخداماً
                if len(self.meta) >= CACHE_MAX_SIZE:
                    self._evict_least_used()

                q_emb = _get_embeddings([question])
                self.index.add(q_emb)
                self.meta.append({
                    "question": question,
                    "answer":   answer,
                    "hits":     0,
                })
                self._save()
                print(f"[CACHE SET] حُفظت إجابة جديدة | إجمالي: {len(self.meta)}")
        except Exception as e:
            print(f"[CACHE SET ERROR] {e}")

    def _evict_least_used(self):
        """احذف أقل 50 إجابة استخداماً وأعد بناء الفهرس"""
        if len(self.meta) < 50:
            return
        # رتّب حسب الاستخدام واحتفظ بالأكثر استخداماً
        self.meta.sort(key=lambda x: x.get("hits", 0), reverse=True)
        keep = self.meta[:CACHE_MAX_SIZE - 50]
        texts = [m["question"] for m in keep]
        embeddings = _get_embeddings(texts)
        new_index = faiss.IndexFlatIP(EMBED_DIM)
        new_index.add(embeddings)
        self.index = new_index
        self.meta  = keep
        print(f"[CACHE EVICT] تم تقليص الذاكرة إلى {len(keep)} إجابة")

    def stats(self):
        return {
            "total":    len(self.meta),
            "max_size": CACHE_MAX_SIZE,
            "threshold": CACHE_SIMILARITY_THRESHOLD,
        }


# ─── instance عام للـ cache ───
semantic_cache = SemanticCache()


# =========================================================
# استخراج النص
# =========================================================

def extract_text(file_path: str, ext: str) -> str:
    ext = ext.lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp"):
        return _extract_image(file_path)
    elif ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"صيغة غير مدعومة: {ext}")


def _extract_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    return "\n".join(
        (p.extract_text() or "") for p in reader.pages
        if (p.extract_text() or "").strip()
    )


def _extract_docx(file_path: str) -> str:
    doc   = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_image(file_path: str) -> str:
    image = Image.open(file_path)
    try:
        ocr_text = pytesseract.image_to_string(image, lang="ara+eng")
    except Exception:
        try:
            ocr_text = pytesseract.image_to_string(image)
        except Exception:
            ocr_text = ""

    ocr_text = ocr_text.strip()
    if len(ocr_text) > 20:
        return ocr_text

    vision_description = _describe_image_with_vision(file_path)
    combined = "\n".join(filter(None, [ocr_text, vision_description]))
    return combined.strip()


def _describe_image_with_vision(file_path: str) -> str:
    if not GROQ_API_KEY:
        return "[لم يتم ضبط GROQ_API_KEY، تعذر تحليل الصورة]"
    try:
        with open(file_path, "rb") as f:
            b64_image = base64.b64encode(f.read()).decode("utf-8")
        ext  = os.path.splitext(file_path)[1].lower().replace(".", "")
        mime = "jpeg" if ext in ("jpg", "jpeg") else ext

        prompt_messages = [{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "صف هذه الصورة بالتفصيل باللغة العربية: "
                        "ما الذي تظهره؟ الأشخاص، الأشياء، النصوص إن وجدت، "
                        "الألوان، السياق العام. اجعل الوصف غنياً بالمعلومات، "
                        "لكن اختصر في حدود 200 كلمة تقريباً حتى لا يُقطع الرد."
                    ),
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/{mime};base64,{b64_image}"},
                },
            ],
        }]

        def _call(reasoning_effort=None):
            kwargs = dict(
                model=VISION_MODEL,
                messages=prompt_messages,
                temperature=0.3,
                # 800 لم تكن كافية: النص العربي التفصيلي يستهلك توكنز أكثر
                # من الإنجليزي لنفس عدد الكلمات، فكان الوصف يُقطع في المنتصف.
                max_completion_tokens=1600,
                # qwen/qwen3.6-27b نموذج تفكير (reasoning)؛ بدون هذا الخيار
                # قد يُرجع سلسلة تفكيره الداخلية بالإنجليزية ضمن content
                # بدل الوصف النهائي فقط، أو يترك content فارغاً في حالات نادرة.
                reasoning_format="hidden",
            )
            if reasoning_effort:
                kwargs["reasoning_effort"] = reasoning_effort
            completion = _groq_client.chat.completions.create(**kwargs)
            text = (completion.choices[0].message.content or "").strip()
            # حماية إضافية: لو تسرّب أي جزء من التفكير رغم reasoning_format=hidden
            text = re.sub(
                r"<think>.*?</think>",
                "",
                text,
                flags=re.DOTALL | re.IGNORECASE,
            ).strip()
            return text

        # المحاولة الأولى: وضع non-thinking (أسرع وأقل عرضة لمشاكل التنسيق)
        description = _call(reasoning_effort="none")

        # لو رجعت فارغة لأي سبب، نعيد المحاولة بالإعداد الافتراضي للنموذج
        if not description:
            description = _call(reasoning_effort=None)

        if not description:
            return "[لم يتمكن نموذج الرؤية من توليد وصف لهذه الصورة]"

        return description
    except Exception as e:
        return f"[تعذر توليد وصف للصورة: {e}]"


# =========================================================
# تقسيم النص
# =========================================================

def chunk_text(text: str, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    text = " ".join(text.split())
    if not text:
        return []
    chunks, start = [], 0
    while start < len(text):
        chunks.append(text[start:start + chunk_size])
        start += chunk_size - overlap
    return chunks


# =========================================================
# قاعدة المعرفة
# =========================================================

class KnowledgeBase:
    def __init__(self):
        self.index = None
        self.meta  = []
        self._load()

    def _load(self):
        if os.path.exists(INDEX_PATH) and os.path.exists(META_PATH):
            self.index = faiss.read_index(INDEX_PATH)
            with open(META_PATH, "r", encoding="utf-8") as f:
                self.meta = json.load(f)
        else:
            self.index = faiss.IndexFlatIP(EMBED_DIM)
            self.meta  = []

    def _save(self):
        faiss.write_index(self.index, INDEX_PATH)
        with open(META_PATH, "w", encoding="utf-8") as f:
            json.dump(self.meta, f, ensure_ascii=False, indent=2)

    def add_document(self, text: str, source_name: str, doc_id: str):
        chunks = chunk_text(text)
        if not chunks:
            return 0
        embeddings = _get_embeddings(chunks)
        with _lock:
            self.index.add(embeddings)
            for i, chunk in enumerate(chunks):
                self.meta.append({
                    "doc_id": doc_id, "source": source_name,
                    "chunk_no": i, "text": chunk,
                })
            self._save()
        return len(chunks)

    def search(self, query: str, top_k=TOP_K):
        if self.index.ntotal == 0:
            return []
        q_emb           = _get_embeddings([query])
        scores, indices = self.index.search(q_emb, min(top_k, self.index.ntotal))
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx == -1:
                continue
            item          = dict(self.meta[idx])
            item["score"] = float(score)
            results.append(item)
        return results

    def list_documents(self):
        seen = {}
        for item in self.meta:
            doc_id = item["doc_id"]
            if doc_id not in seen:
                seen[doc_id] = {"doc_id": doc_id, "source": item["source"], "chunks": 0}
            seen[doc_id]["chunks"] += 1
        return list(seen.values())

    def delete_document(self, doc_id: str):
        keep_meta = [m for m in self.meta if m["doc_id"] != doc_id]
        if len(keep_meta) == len(self.meta):
            return False
        with _lock:
            if keep_meta:
                new_index = faiss.IndexFlatIP(EMBED_DIM)
                new_index.add(_get_embeddings([m["text"] for m in keep_meta]))
            else:
                new_index = faiss.IndexFlatIP(EMBED_DIM)
            self.index = new_index
            self.meta  = keep_meta
            self._save()
        return True


def save_uploaded_file(file_storage, original_filename: str):
    ext       = os.path.splitext(original_filename)[1].lower()
    doc_id    = uuid.uuid4().hex[:12]
    file_path = os.path.join(UPLOADS_DIR, f"{doc_id}{ext}")
    file_storage.save(file_path)
    return doc_id, file_path, ext



def describe_uploaded_image(filename: str) -> str:
    """
    يصف صورة مرفوعة مسبقاً في قاعدة المعرفة باستخدام Vision.
    يبحث عن الصورة في مجلد uploads ويعيد وصفها.
    """
    # ابحث عن الملف في مجلد uploads
    if not filename:
        return None

    # البحث بالاسم الأصلي أو بالـ doc_id
    target_path = None

    # أولاً: بحث مباشر
    direct = os.path.join(UPLOADS_DIR, filename)
    if os.path.exists(direct):
        target_path = direct

    # ثانياً: بحث في المتا عبر اسم الملف
    if not target_path:
        for item in kb_meta_search(filename):
            candidate = os.path.join(UPLOADS_DIR, f"{item['doc_id']}{os.path.splitext(filename)[1]}")
            if os.path.exists(candidate):
                target_path = candidate
                break

    # ثالثاً: أخذ آخر صورة مرفوعة
    if not target_path:
        image_exts = {'.jpg', '.jpeg', '.png', '.webp', '.bmp'}
        candidates = [
            os.path.join(UPLOADS_DIR, f)
            for f in os.listdir(UPLOADS_DIR)
            if os.path.splitext(f)[1].lower() in image_exts
        ]
        if candidates:
            target_path = max(candidates, key=os.path.getmtime)

    if not target_path:
        return None

    return _describe_image_with_vision(target_path)


def kb_meta_search(source_name: str) -> list:
    """بحث في الـ meta عن اسم مصدر معين"""
    if not os.path.exists(META_PATH):
        return []
    try:
        with open(META_PATH, 'r', encoding='utf-8') as f:
            meta = json.load(f)
        return [m for m in meta if source_name.lower() in m.get('source', '').lower()]
    except Exception:
        return []


def get_last_uploaded_image_path() -> str:
    """يُعيد مسار آخر صورة مرفوعة"""
    image_exts = {'.jpg', '.jpeg', '.png', '.webp', '.bmp'}
    try:
        candidates = [
            os.path.join(UPLOADS_DIR, f)
            for f in os.listdir(UPLOADS_DIR)
            if os.path.splitext(f)[1].lower() in image_exts
        ]
        if candidates:
            return max(candidates, key=os.path.getmtime)
    except Exception:
        pass
    return None


def build_context_block(results):
    if not results:
        return None
    lines = ["## معلومات ذات صلة من المستندات المرفوعة:\n"]
    for r in results:
        lines.append(f"[من: {r['source']}]\n{r['text']}\n")
    return "\n".join(lines)
